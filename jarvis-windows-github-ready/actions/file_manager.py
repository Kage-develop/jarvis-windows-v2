"""
Dosya yönetimi — TXT oluştur/yaz/oku, DOCX oluştur/yaz, sil/taşı/kopyala
"""

from __future__ import annotations

import os
import shutil
import time
from pathlib import Path

# python-docx kütüphanesi (pip install python-docx)
try:
    from docx import Document as DocxDocument
    from docx.shared import Pt
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


# ---------------------------------------------------------------------------
# Yardımcı: Güvenli path çözümleme
# ---------------------------------------------------------------------------

def _resolve_path(path: str) -> Path:
    """
    Verilen path'i çöz:
    - ~ → kullanıcı home dizini
    - Göreceli path → kullanıcının Desktop'u
    - Mutlak path → olduğu gibi
    """
    p = Path(os.path.expandvars(os.path.expanduser(path.strip())))
    if not p.is_absolute():
        # Göreceli path'i Desktop'a göre çöz
        desktop = Path.home() / "Desktop"
        p = desktop / p
    return p


def _write_roots() -> list[Path]:
    home = Path.home().resolve()
    roots = [
        home,
        home / "Desktop",
        home / "Documents",
        home / "Downloads",
        home / "OneDrive",
    ]
    seen: set[str] = set()
    unique: list[Path] = []
    for root in roots:
        try:
            resolved = root.resolve()
        except OSError:
            continue
        key = str(resolved).casefold()
        if key not in seen:
            seen.add(key)
            unique.append(resolved)
    return unique


def _safe_check(path: Path) -> str | None:
    """Tehlikeli dizin + izinli kok disi silme/yazma engeli."""
    blocked = ["system32", "windows", "program files", "programdata"]
    for part in path.parts:
        if part.lower() in blocked:
            return f"Güvenlik: Bu dizine erisim engellendi → {path}"

    try:
        resolved = path.resolve()
    except OSError:
        return f"Güvenlik: Gecersiz yol → {path}"

    for root in _write_roots():
        try:
            resolved.relative_to(root)
            return None
        except ValueError:
            continue

    allowed = ", ".join(str(r) for r in _write_roots())
    return (
        f"Güvenlik: Bu yol izinli degil → {path}. "
        f"Izinli: {allowed}"
    )


# ---------------------------------------------------------------------------
# TXT işlemleri
# ---------------------------------------------------------------------------

def write_text_file(path: str, content: str, mode: str = "w") -> str:
    """
    TXT dosyasına yaz.

    Args:
        path:    Dosya yolu (göreceli → Desktop'a göre çözülür)
        content: Yazılacak içerik
        mode:    'w' = üzerine yaz (varsayılan), 'a' = ekle (append)
    """
    p = _resolve_path(path)
    err = _safe_check(p)
    if err:
        return err

    p.parent.mkdir(parents=True, exist_ok=True)

    try:
        with open(p, mode, encoding="utf-8") as f:
            f.write(content)
        action = "eklendi" if mode == "a" else "yazıldı"
        return f"✓ {p.name} dosyasına içerik {action}. ({p})"
    except Exception as e:
        return f"Hata: {e}"


def read_text_file(path: str) -> str:
    """TXT dosyasını oku ve içeriğini döndür."""
    p = _resolve_path(path)

    if not p.exists():
        return f"Dosya bulunamadı: {p}"

    try:
        content = p.read_text(encoding="utf-8", errors="replace")
        if len(content) > 2000:
            content = content[:2000] + "\n... (içerik kısaltıldı)"
        return content or "(Dosya boş)"
    except Exception as e:
        return f"Hata: {e}"


def create_text_file(path: str, content: str = "") -> str:
    """Yeni TXT dosyası oluştur (zaten varsa içeriğini korur)."""
    p = _resolve_path(path)

    if p.exists():
        return f"Dosya zaten mevcut: {p}"

    return write_text_file(path, content, mode="w")


# ---------------------------------------------------------------------------
# DOCX işlemleri
# ---------------------------------------------------------------------------

def create_docx_file(path: str, title: str = "", content: str = "") -> str:
    """
    Yeni Word belgesi (.docx) oluştur.

    Args:
        path:    Dosya yolu (.docx uzantısıyla)
        title:   Belge başlığı (isteğe bağlı)
        content: Belge içeriği — satır satır paragraf olarak eklenir
    """
    if not DOCX_AVAILABLE:
        return (
            "python-docx kütüphanesi yüklü değil. "
            "Yüklemek için: pip install python-docx"
        )

    if not path.endswith(".docx"):
        path += ".docx"

    p = _resolve_path(path)
    err = _safe_check(p)
    if err:
        return err

    p.parent.mkdir(parents=True, exist_ok=True)

    try:
        doc = DocxDocument()

        if title:
            heading = doc.add_heading(title, level=1)
            heading.runs[0].font.size = Pt(16)

        if content:
            for line in content.splitlines():
                doc.add_paragraph(line)

        doc.save(str(p))
        return f"✓ Word belgesi oluşturuldu: {p}"
    except Exception as e:
        return f"Hata: {e}"


def append_to_docx(path: str, content: str, heading: str = "") -> str:
    """
    Mevcut .docx dosyasına içerik ekle.

    Args:
        path:    Dosya yolu
        content: Eklenecek metin
        heading: İsteğe bağlı bölüm başlığı
    """
    if not DOCX_AVAILABLE:
        return "python-docx kütüphanesi yüklü değil. Yüklemek için: pip install python-docx"

    if not path.endswith(".docx"):
        path += ".docx"

    p = _resolve_path(path)

    if not p.exists():
        return create_docx_file(path, heading, content)

    try:
        doc = DocxDocument(str(p))

        if heading:
            doc.add_heading(heading, level=2)

        for line in content.splitlines():
            doc.add_paragraph(line)

        doc.save(str(p))
        return f"✓ {p.name} dosyasına içerik eklendi."
    except Exception as e:
        return f"Hata: {e}"


# ---------------------------------------------------------------------------
# Dosya yönetimi
# ---------------------------------------------------------------------------

def delete_file(path: str) -> str:
    """Dosya veya boş klasör sil."""
    p = _resolve_path(path)
    err = _safe_check(p)
    if err:
        return err

    if not p.exists():
        return f"Dosya/klasör bulunamadı: {p}"

    try:
        if p.is_file():
            p.unlink()
            return f"✓ Silindi: {p}"
        elif p.is_dir():
            if any(p.iterdir()):
                return f"Klasör boş değil, silmek için önce içini temizle: {p}"
            p.rmdir()
            return f"✓ Klasör silindi: {p}"
    except Exception as e:
        return f"Hata: {e}"


def move_file(src: str, dst: str) -> str:
    """Dosyayı/klasörü taşı veya yeniden adlandır."""
    src_p = _resolve_path(src)
    dst_p = _resolve_path(dst)

    err = _safe_check(dst_p)
    if err:
        return err

    if not src_p.exists():
        return f"Kaynak bulunamadı: {src_p}"

    try:
        dst_p.parent.mkdir(parents=True, exist_ok=True)
        shutil.move(str(src_p), str(dst_p))
        return f"✓ Taşındı: {src_p.name} → {dst_p}"
    except Exception as e:
        return f"Hata: {e}"


def copy_file(src: str, dst: str) -> str:
    """Dosyayı kopyala."""
    src_p = _resolve_path(src)
    dst_p = _resolve_path(dst)

    err = _safe_check(dst_p)
    if err:
        return err

    if not src_p.exists():
        return f"Kaynak bulunamadı: {src_p}"

    try:
        dst_p.parent.mkdir(parents=True, exist_ok=True)
        if src_p.is_dir():
            shutil.copytree(str(src_p), str(dst_p))
        else:
            shutil.copy2(str(src_p), str(dst_p))
        return f"✓ Kopyalandı: {src_p.name} → {dst_p}"
    except Exception as e:
        return f"Hata: {e}"


def list_files(path: str = "~/Desktop") -> str:
    """Klasördeki dosyaları listele."""
    p = _resolve_path(path)

    if not p.exists():
        return f"Klasör bulunamadı: {p}"

    try:
        items = sorted(p.iterdir(), key=lambda x: (x.is_file(), x.name.lower()))
        if not items:
            return f"Klasör boş: {p}"

        lines = [f"📁 {p}", ""]
        for item in items[:50]:  # max 50 dosya göster
            icon = "📄" if item.is_file() else "📁"
            lines.append(f"  {icon} {item.name}")

        if len(list(p.iterdir())) > 50:
            lines.append("  ... (daha fazlası var)")

        return "\n".join(lines)
    except Exception as e:
        return f"Hata: {e}"


def _format_size(size: int) -> str:
    units = ("B", "KB", "MB", "GB")
    value = float(size)
    for unit in units:
        if value < 1024 or unit == units[-1]:
            return f"{value:.1f}{unit}" if unit != "B" else f"{int(value)}B"
        value /= 1024
    return f"{size}B"


def _safe_scan_root(path: str) -> Path:
    p = _resolve_path(path or "~/Desktop")
    home = Path.home().resolve()
    try:
        resolved = p.resolve()
        resolved.relative_to(home)
        return resolved
    except (OSError, ValueError):
        return home / "Desktop"


def _iter_files(root: Path, max_seen: int = 5000):
    seen = 0
    for base, dirs, files in os.walk(root):
        dirs[:] = [
            d for d in dirs
            if d.lower() not in {"__pycache__", ".git", "node_modules", ".venv", "venv"}
        ]
        for name in files:
            path = Path(base) / name
            try:
                path.stat()
            except OSError:
                continue
            yield path
            seen += 1
            if seen >= max_seen:
                return


def find_files(query: str = "", path: str = "~/Desktop", extension: str = "", limit: int = 20) -> str:
    """Dosya adina gore guvenli arama yap."""
    root = _safe_scan_root(path)
    q = (query or "").strip().lower()
    ext = (extension or "").strip().lower()
    if ext and not ext.startswith("."):
        ext = "." + ext
    limit = max(1, min(80, int(limit or 20)))

    matches = []
    for item in _iter_files(root):
        name = item.name.lower()
        if q and q not in name:
            continue
        if ext and item.suffix.lower() != ext:
            continue
        try:
            st = item.stat()
        except OSError:
            continue
        matches.append((st.st_mtime, st.st_size, item))
        if len(matches) >= limit:
            break

    if not matches:
        return f"Eslesen dosya bulunamadi: {root}"

    matches.sort(reverse=True)
    lines = [f"Eslesen dosyalar ({root}):"]
    for mtime, size, item in matches[:limit]:
        age_days = max(0, int((time.time() - mtime) // 86400))
        lines.append(f"- {item.name} | {_format_size(size)} | {age_days} gun once | {item}")
    return "\n".join(lines)


def recent_files(path: str = "~/Downloads", limit: int = 15) -> str:
    """Klasordeki son degisen dosyalari listele."""
    root = _safe_scan_root(path)
    limit = max(1, min(50, int(limit or 15)))
    files = []
    for item in _iter_files(root):
        try:
            st = item.stat()
        except OSError:
            continue
        files.append((st.st_mtime, st.st_size, item))
    files.sort(reverse=True)
    if not files:
        return f"Dosya bulunamadi: {root}"
    lines = [f"Son dosyalar ({root}):"]
    for mtime, size, item in files[:limit]:
        stamp = time.strftime("%Y-%m-%d %H:%M", time.localtime(mtime))
        lines.append(f"- {stamp} | {_format_size(size)} | {item.name} | {item}")
    return "\n".join(lines)


def largest_files(path: str = "~/Downloads", limit: int = 15, min_mb: int = 10) -> str:
    """Buyuk dosyalari listele."""
    root = _safe_scan_root(path)
    limit = max(1, min(50, int(limit or 15)))
    threshold = max(0, int(min_mb or 0)) * 1024 * 1024
    files = []
    for item in _iter_files(root):
        try:
            size = item.stat().st_size
        except OSError:
            continue
        if size >= threshold:
            files.append((size, item))
    files.sort(reverse=True)
    if not files:
        return f"{root} icinde {min_mb}MB ustu dosya bulunamadi."
    lines = [f"Buyuk dosyalar ({root}):"]
    for size, item in files[:limit]:
        lines.append(f"- {_format_size(size)} | {item.name} | {item}")
    return "\n".join(lines)


def downloads_summary(limit: int = 10) -> str:
    """Downloads klasoru icin hizli ozet."""
    root = Path.home() / "Downloads"
    if not root.exists():
        return "Downloads klasoru bulunamadi."
    by_ext: dict[str, tuple[int, int]] = {}
    total = 0
    count = 0
    for item in _iter_files(root):
        try:
            size = item.stat().st_size
        except OSError:
            continue
        ext = item.suffix.lower() or "(uzantisiz)"
        old_count, old_size = by_ext.get(ext, (0, 0))
        by_ext[ext] = (old_count + 1, old_size + size)
        total += size
        count += 1
    top = sorted(by_ext.items(), key=lambda kv: kv[1][1], reverse=True)[:max(3, min(20, limit))]
    lines = [f"Downloads ozeti: {count} dosya, toplam {_format_size(total)}"]
    for ext, (ext_count, ext_size) in top:
        lines.append(f"- {ext}: {ext_count} dosya, {_format_size(ext_size)}")
    return "\n".join(lines)


def organize_downloads(dry_run: bool = True) -> str:
    """Downloads dosyalarini tip klasorlerine ayir. dry_run=True iken sadece plan verir."""
    root = Path.home() / "Downloads"
    if not root.exists():
        return "Downloads klasoru bulunamadi."

    groups = {
        "Images": {".png", ".jpg", ".jpeg", ".webp", ".gif", ".bmp"},
        "Documents": {".pdf", ".doc", ".docx", ".txt", ".rtf", ".xlsx", ".pptx", ".csv"},
        "Archives": {".zip", ".rar", ".7z", ".tar", ".gz"},
        "Installers": {".exe", ".msi", ".msix"},
        "Media": {".mp3", ".wav", ".mp4", ".mov", ".mkv", ".avi"},
    }
    moves = []
    for item in root.iterdir():
        if not item.is_file():
            continue
        target_group = None
        for group, exts in groups.items():
            if item.suffix.lower() in exts:
                target_group = group
                break
        if not target_group:
            continue
        dst = root / target_group / item.name
        moves.append((item, dst))

    if not moves:
        return "Downloads zaten duzenli gorunuyor."

    if dry_run:
        lines = ["Downloads duzenleme plani:"]
        for src, dst in moves[:40]:
            lines.append(f"- {src.name} -> {dst.parent.name}/")
        if len(moves) > 40:
            lines.append(f"... {len(moves) - 40} dosya daha")
        lines.append("Uygulamak icin organize_downloads dry_run=false kullan.")
        return "\n".join(lines)

    for src, dst in moves:
        try:
            dst.parent.mkdir(parents=True, exist_ok=True)
            if dst.exists():
                dst = dst.with_name(f"{dst.stem}-{int(time.time())}{dst.suffix}")
            shutil.move(str(src), str(dst))
        except Exception:
            continue
    return f"Downloads duzenlendi: {len(moves)} dosya tasindi."


def file_assistant(
    action: str,
    path: str = "~/Desktop",
    query: str = "",
    extension: str = "",
    limit: int = 20,
    min_mb: int = 10,
    dry_run: bool = True,
) -> str:
    """Dosya asistani tek giris noktasi."""
    act = (action or "").strip().lower().replace(" ", "_")
    if act in {"search", "find", "ara", "bul"}:
        return find_files(query, path, extension, limit)
    if act in {"recent", "last", "son", "son_dosyalar"}:
        return recent_files(path, limit)
    if act in {"largest", "big", "buyuk", "buyuk_dosyalar"}:
        return largest_files(path, limit, min_mb)
    if act in {"downloads_summary", "indirilenler_ozet", "downloads"}:
        return downloads_summary(limit)
    if act in {"organize_downloads", "indirilenleri_duzenle"}:
        return organize_downloads(dry_run)
    return "file_assistant action: search | recent | largest | downloads_summary | organize_downloads"
